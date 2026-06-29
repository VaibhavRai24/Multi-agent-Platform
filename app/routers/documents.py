import uuid
from datetime import datetime

from fastapi import APIRouter, UploadFile, File, HTTPException, Depends
from sqlalchemy.orm import Session

from app.database.database import get_db
from app.models.domain import Document
from app.schemas.schemas import DocumentResponse
from app.utils.logger import get_logger
from app.utils.pinecone_client import get_pinecone_index, embed_texts

logger = get_logger(__name__)
router = APIRouter()

CHUNK_SIZE = 800
CHUNK_OVERLAP = 100

SUPPORTED_TYPES = {
    ".pdf", ".docx", ".txt", ".md", ".csv",
    ".png", ".jpg", ".jpeg", ".webp", ".gif",
    ".xlsx", ".xls",
}


def get_file_ext(filename: str) -> str:
    return "." + filename.rsplit(".", 1)[-1].lower() if "." in filename else ""


def extract_text(filename: str, content: bytes) -> str:
    """Extract raw text from an uploaded file based on extension."""
    ext = get_file_ext(filename)

    if ext == ".pdf":
        import io
        from pypdf import PdfReader
        reader = PdfReader(io.BytesIO(content))
        return "\n".join(page.extract_text() or "" for page in reader.pages)

    if ext == ".docx":
        import io
        try:
            from docx import Document as DocxDocument
            doc = DocxDocument(io.BytesIO(content))
            paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
            # Also extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            paragraphs.append(cell.text.strip())
            return "\n".join(paragraphs)
        except ImportError:
            # Fallback to docx2txt with temp file if python-docx not installed
            import tempfile, os
            import docx2txt
            with tempfile.NamedTemporaryFile(delete=False, suffix=".docx") as tmp:
                tmp.write(content)
                tmp_path = tmp.name
            try:
                text = docx2txt.process(tmp_path) or ""
            finally:
                os.unlink(tmp_path)
            return text

    if ext in (".txt", ".md", ".csv"):
        return content.decode("utf-8", errors="ignore")

    if ext in (".xlsx", ".xls"):
        import io
        import openpyxl
        wb = openpyxl.load_workbook(io.BytesIO(content), read_only=True, data_only=True)
        lines = []
        for ws in wb.worksheets:
            for row in ws.iter_rows(values_only=True):
                line = "\t".join(str(c) if c is not None else "" for c in row)
                if line.strip():
                    lines.append(line)
        return "\n".join(lines)

    if ext in (".png", ".jpg", ".jpeg", ".webp", ".gif"):
        # For images, we store a descriptive placeholder so the LLM
        # at least knows the image exists in the knowledge base.
        return f"[Image file: {filename}] — This document is an image uploaded by the user to the knowledge base."

    raise HTTPException(
        status_code=400,
        detail=f"Unsupported file type '{ext}'. Supported: pdf, docx, txt, md, csv, xlsx, png, jpg, jpeg, webp, gif.",
    )


def chunk_text(text: str, chunk_size: int = CHUNK_SIZE, overlap: int = CHUNK_OVERLAP) -> list[str]:
    """Split text into overlapping chunks for embedding."""
    text = text.strip()
    if not text:
        return []
    chunks = []
    start = 0
    while start < len(text):
        end = start + chunk_size
        chunks.append(text[start:end])
        start += chunk_size - overlap
    return [c.strip() for c in chunks if c.strip()]


# ── Hardcoded user_id=1 until JWT auth is wired to document endpoints ──
MOCK_USER_ID = 1


@router.post("/upload", response_model=DocumentResponse)
def upload_document(file: UploadFile = File(...), db: Session = Depends(get_db)):
    ext = get_file_ext(file.filename)
    if ext not in SUPPORTED_TYPES:
        raise HTTPException(status_code=400, detail=f"Unsupported file type '{ext}'.")

    content = file.file.read()

    try:
        text = extract_text(file.filename, content)
    except HTTPException:
        raise
    except Exception as exc:
        logger.error(f"Text extraction failed for {file.filename}: {exc}")
        raise HTTPException(status_code=400, detail=f"Failed to parse file: {exc}")

    chunks = chunk_text(text)
    document_id = str(uuid.uuid4())
    file_type = ext.lstrip(".")
    chunks_indexed = 0
    status = "processed"

    # --- Pinecone upsert (if configured) ---
    index = get_pinecone_index()
    if index is not None and chunks:
        try:
            vectors = embed_texts(chunks)
            upserts = [
                {
                    "id": f"{document_id}-{i}",
                    "values": vector,
                    "metadata": {
                        "text": chunk,
                        "source": file.filename,
                        "document_id": document_id,
                        "user_id": MOCK_USER_ID,
                        "chunk_index": i,
                    },
                }
                for i, (chunk, vector) in enumerate(zip(chunks, vectors))
            ]
            index.upsert(vectors=upserts)
            chunks_indexed = len(chunks)
            logger.info(f"Indexed {chunks_indexed} chunks for '{file.filename}' (doc_id={document_id})")
        except Exception as exc:
            logger.error(f"Pinecone upsert failed for {file.filename}: {exc}")
            status = "failed"
    elif not chunks:
        status = "failed"
        logger.warning(f"No extractable text in '{file.filename}'.")
    else:
        logger.warning("Pinecone not configured — chunks not indexed, document saved as reference only.")
        chunks_indexed = 0
        status = "no_index"

    # --- Persist to SQLite ---
    doc = Document(
        title=file.filename,
        document_id=document_id,
        file_type=file_type,
        status=status,
        chunks_indexed=chunks_indexed,
        user_id=MOCK_USER_ID,
    )
    db.add(doc)
    db.commit()
    db.refresh(doc)

    return doc


@router.get("/", response_model=list[DocumentResponse])
def list_documents(db: Session = Depends(get_db)):
    docs = db.query(Document).filter(Document.user_id == MOCK_USER_ID).order_by(Document.created_at.desc()).all()
    return docs


@router.delete("/{doc_id}")
def delete_document(doc_id: int, db: Session = Depends(get_db)):
    doc = db.query(Document).filter(Document.id == doc_id, Document.user_id == MOCK_USER_ID).first()
    if not doc:
        raise HTTPException(status_code=404, detail="Document not found.")

    # Remove vectors from Pinecone
    index = get_pinecone_index()
    if index is not None and doc.chunks_indexed > 0:
        try:
            ids_to_delete = [f"{doc.document_id}-{i}" for i in range(doc.chunks_indexed)]
            index.delete(ids=ids_to_delete)
            logger.info(f"Deleted {len(ids_to_delete)} vectors for doc_id={doc.document_id}")
        except Exception as exc:
            logger.warning(f"Pinecone vector deletion failed: {exc}")

    db.delete(doc)
    db.commit()
    return {"status": "deleted", "id": doc_id}
