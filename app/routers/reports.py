from fastapi import APIRouter, Depends, HTTPException
from fastapi.responses import StreamingResponse
from sqlalchemy.orm import Session
from io import BytesIO
import subprocess, tempfile, os, re

from app.database.database import get_db
from app.models.domain import Report
from app.utils.logger import get_logger

logger = get_logger(__name__)
router = APIRouter()

MOCK_USER_ID = 1


@router.get("/")
def list_reports(db: Session = Depends(get_db)):
    reports = (
        db.query(Report)
        .filter(Report.user_id == MOCK_USER_ID)
        .order_by(Report.created_at.desc())
        .all()
    )
    return [
        {
            "id": r.id,
            "title": r.title,
            "format": r.content_format,
            "preview": r.content[:300] if r.content else "",
            "date": r.created_at.strftime("%d %b %Y, %H:%M") if r.created_at else "",
        }
        for r in reports
    ]


@router.post("/save")
def save_report(payload: dict, db: Session = Depends(get_db)):
    title = payload.get("title", "Untitled Report")
    content = payload.get("content", "")
    fmt = payload.get("format", "markdown")

    report = Report(
        title=title,
        content_format=fmt,
        content=content,
        user_id=MOCK_USER_ID,
    )
    db.add(report)
    db.commit()
    db.refresh(report)
    logger.info(f"Saved report #{report.id}: {title}")
    return {"id": report.id, "title": report.title, "status": "saved"}


def _markdown_to_docx_bytes(title: str, markdown_content: str) -> bytes:
    """Convert markdown text to a .docx file using python-docx."""
    try:
        from docx import Document
        from docx.shared import Pt, RGBColor, Inches
        from docx.enum.text import WD_ALIGN_PARAGRAPH
    except ImportError:
        raise HTTPException(
            status_code=500,
            detail="python-docx is not installed. Run: pip install python-docx"
        )

    doc = Document()

    # Title page heading
    title_para = doc.add_heading(title, level=0)
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()  # spacer

    lines = markdown_content.split("\n")
    i = 0
    while i < len(lines):
        line = lines[i]

        # Headings
        if line.startswith("### "):
            doc.add_heading(line[4:].strip(), level=3)
        elif line.startswith("## "):
            doc.add_heading(line[3:].strip(), level=2)
        elif line.startswith("# "):
            doc.add_heading(line[2:].strip(), level=1)

        # Horizontal rule → spacer paragraph
        elif line.strip() in ("---", "***", "___"):
            p = doc.add_paragraph()
            pPr = p._p.get_or_add_pPr()
            from docx.oxml.ns import qn
            from lxml import etree
            border_el = etree.SubElement(pPr, qn("w:pBdr"))
            bot = etree.SubElement(border_el, qn("w:bottom"))
            bot.set(qn("w:val"), "single")
            bot.set(qn("w:sz"), "6")
            bot.set(qn("w:space"), "1")
            bot.set(qn("w:color"), "AAAAAA")

        # Bullet / unordered list
        elif re.match(r"^[-*+] ", line):
            doc.add_paragraph(line[2:].strip(), style="List Bullet")

        # Numbered list
        elif re.match(r"^\d+\. ", line):
            doc.add_paragraph(re.sub(r"^\d+\. ", "", line).strip(), style="List Number")

        # Code block
        elif line.startswith("```"):
            code_lines = []
            i += 1
            while i < len(lines) and not lines[i].startswith("```"):
                code_lines.append(lines[i])
                i += 1
            code_para = doc.add_paragraph("\n".join(code_lines))
            code_para.style = doc.styles["No Spacing"]
            run = code_para.runs[0] if code_para.runs else code_para.add_run()
            run.font.name = "Courier New"
            run.font.size = Pt(9)

        # Bold / italic inline in normal paragraph
        elif line.strip():
            para = doc.add_paragraph()
            # Simple inline bold (**text**) and italic (*text*) parsing
            remaining = line
            pattern = re.compile(r"(\*\*(.+?)\*\*|\*(.+?)\*|`(.+?)`)")
            last = 0
            for m in pattern.finditer(remaining):
                # plain text before match
                if m.start() > last:
                    para.add_run(remaining[last:m.start()])
                full = m.group(0)
                if full.startswith("**"):
                    run = para.add_run(m.group(2))
                    run.bold = True
                elif full.startswith("`"):
                    run = para.add_run(m.group(4))
                    run.font.name = "Courier New"
                    run.font.size = Pt(9)
                else:
                    run = para.add_run(m.group(3))
                    run.italic = True
                last = m.end()
            if last < len(remaining):
                para.add_run(remaining[last:])

        else:
            # blank line → spacer
            doc.add_paragraph()

        i += 1

    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()


@router.get("/download/{report_id}")
def download_report(report_id: int, fmt: str = "markdown", db: Session = Depends(get_db)):
    """Download a report as markdown (.md) or Word (.docx).

    Query param:  ?fmt=markdown  (default)  |  ?fmt=docx
    """
    report = db.query(Report).filter(
        Report.id == report_id, Report.user_id == MOCK_USER_ID
    ).first()
    if not report:
        raise HTTPException(status_code=404, detail="Report not found.")

    safe_title = report.title.replace(" ", "_").replace("/", "-")[:60]

    if fmt == "docx":
        docx_bytes = _markdown_to_docx_bytes(report.title, report.content)
        return StreamingResponse(
            BytesIO(docx_bytes),
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={"Content-Disposition": f'attachment; filename="{safe_title}.docx"'},
        )

    # default: markdown
    content_bytes = report.content.encode("utf-8")
    return StreamingResponse(
        BytesIO(content_bytes),
        media_type="text/markdown",
        headers={"Content-Disposition": f'attachment; filename="{safe_title}.md"'},
    )


@router.get("/view/{report_id}")
def view_report(report_id: int, db: Session = Depends(get_db)):
    report = db.query(Report).filter(
        Report.id == report_id, Report.user_id == MOCK_USER_ID
    ).first()
    if not report:
        raise HTTPException(status_code=404, detail="Report not found.")
    return {
        "id": report.id,
        "title": report.title,
        "content": report.content,
        "format": report.content_format,
        "date": report.created_at.strftime("%d %b %Y, %H:%M") if report.created_at else "",
    }


@router.delete("/{report_id}")
def delete_report(report_id: int, db: Session = Depends(get_db)):
    report = db.query(Report).filter(
        Report.id == report_id, Report.user_id == MOCK_USER_ID
    ).first()
    if not report:
        raise HTTPException(status_code=404, detail="Report not found.")
    db.delete(report)
    db.commit()
    return {"status": "deleted", "id": report_id}
